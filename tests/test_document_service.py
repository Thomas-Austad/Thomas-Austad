from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.services import document_service
from app.services.document_service import (
    DOCX_CONTENT_TYPE,
    MAX_RESUME_UPLOAD_BYTES,
    PDF_CONTENT_TYPE,
    ResumeExtractionError,
    extract_resume_text,
    markdown_resume_to_docx,
)

LOCAL_AUTH_HEADERS = {
    "Authorization": "Bearer test-local-access-token-that-is-at-least-32-characters"
}


def local_client() -> TestClient:
    return TestClient(app, headers=LOCAL_AUTH_HEADERS)


def test_docx_generation():
    result = markdown_resume_to_docx("# Jane Doe\n## Experience\n- Improved throughput 20%")
    assert result[:2] == b"PK"


def test_docx_resume_extraction():
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Built Python APIs")
    stream = BytesIO()
    doc.save(stream)

    result = extract_resume_text(stream.getvalue(), "resume.docx")

    assert result.filename == "resume.docx"
    assert result.content_type == DOCX_CONTENT_TYPE
    assert result.text == "Jane Doe\nBuilt Python APIs"
    assert result.character_count == len(result.text)


def test_pdf_resume_extraction():
    pdf = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT /F1 24 Tf 100 700 Td (Jane Doe PDF) Tj ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000241 00000 n 
0000000335 00000 n 
trailer
<< /Root 1 0 R /Size 6 >>
startxref
405
%%EOF
"""

    result = extract_resume_text(pdf, "resume.pdf")

    assert result.filename == "resume.pdf"
    assert result.content_type == PDF_CONTENT_TYPE
    assert result.text == "Jane Doe PDF"


def test_resume_extraction_rejects_unsupported_content():
    with pytest.raises(ResumeExtractionError, match="Unsupported resume file type"):
        extract_resume_text(b"not a resume", "resume.txt")


def test_resume_extraction_rejects_zip_that_is_not_docx():
    with pytest.raises(ResumeExtractionError, match="Unsupported resume file type"):
        extract_resume_text(b"PK\x03\x04not a valid docx", "resume.docx")


def test_resume_extraction_rejects_oversized_content():
    with pytest.raises(ResumeExtractionError, match="too large"):
        extract_resume_text(b"x" * (MAX_RESUME_UPLOAD_BYTES + 1), "resume.pdf")


def test_docx_resume_extraction_rejects_excessive_archive_members(monkeypatch):
    monkeypatch.setattr(document_service, "MAX_DOCX_ARCHIVE_MEMBERS", 1)
    content = _docx_archive({"word/document.xml": "<document />", "extra.xml": "<extra />"})

    with pytest.raises(ResumeExtractionError, match="too many files"):
        extract_resume_text(content, "resume.docx")


def test_docx_resume_extraction_rejects_excessive_uncompressed_content(monkeypatch):
    monkeypatch.setattr(document_service, "MAX_DOCX_UNCOMPRESSED_BYTES", 8)
    content = _docx_archive({"word/document.xml": "x" * 9})

    with pytest.raises(ResumeExtractionError, match="too large after decompression"):
        extract_resume_text(content, "resume.docx")


def test_docx_resume_extraction_rejects_excessive_compression_ratio(monkeypatch):
    monkeypatch.setattr(document_service, "MAX_DOCX_COMPRESSION_RATIO", 2)
    content = _docx_archive({"word/document.xml": "x" * 1_000})

    with pytest.raises(ResumeExtractionError, match="safe compression ratio"):
        extract_resume_text(content, "resume.docx")


def test_docx_resume_extraction_rejects_unsafe_archive_path():
    content = _docx_archive({"../word/document.xml": "<document />"})

    with pytest.raises(ResumeExtractionError, match="unsafe file path"):
        extract_resume_text(content, "resume.docx")


def _docx_archive(entries: dict[str, str]) -> bytes:
    stream = BytesIO()
    with ZipFile(stream, "w", ZIP_DEFLATED) as archive:
        for filename, value in entries.items():
            archive.writestr(filename, value)
    return stream.getvalue()


def test_resume_extract_endpoint_accepts_docx():
    content = markdown_resume_to_docx("# Jane Doe\nBuilt Python APIs")

    response = local_client().post(
        "/resumes/extract",
        files={
            "file": (
                "resume.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert response.json()["filename"] == "resume.docx"
    assert "Jane Doe" in response.json()["text"]


def test_resume_extract_endpoint_rejects_unsupported_file():
    response = local_client().post(
        "/resumes/extract",
        files={"file": ("resume.txt", b"plain text", "text/plain")},
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "Unsupported resume file type"


def test_resume_extract_endpoint_rejects_unsafe_docx_archive(monkeypatch):
    monkeypatch.setattr(document_service, "MAX_DOCX_ARCHIVE_MEMBERS", 1)
    content = _docx_archive({"word/document.xml": "<document />", "extra.xml": "<extra />"})

    response = local_client().post(
        "/resumes/extract",
        files={"file": ("resume.docx", content, DOCX_CONTENT_TYPE)},
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "DOCX archive contains too many files"
