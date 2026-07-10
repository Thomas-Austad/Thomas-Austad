from io import BytesIO
from pathlib import PurePosixPath
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader

from app.models.schemas import ExtractedResume

MAX_RESUME_UPLOAD_BYTES = 5 * 1024 * 1024
MAX_DOCX_ARCHIVE_MEMBERS = 256
MAX_DOCX_UNCOMPRESSED_BYTES = 20 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 100
PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class ResumeExtractionError(ValueError):
    pass


def markdown_resume_to_docx(markdown_text: str) -> bytes:
    doc = Document()
    for raw in markdown_text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def extract_resume_text(content: bytes, filename: str | None = None) -> ExtractedResume:
    if len(content) > MAX_RESUME_UPLOAD_BYTES:
        raise ResumeExtractionError("Resume file is too large")
    if content.startswith(b"%PDF-"):
        text = _extract_pdf_text(content)
        content_type = PDF_CONTENT_TYPE
    elif content.startswith(b"PK"):
        text = _extract_docx_text(content)
        content_type = DOCX_CONTENT_TYPE
    else:
        raise ResumeExtractionError("Unsupported resume file type")

    normalized = _normalize_extracted_text(text)
    if not normalized:
        raise ResumeExtractionError("Resume file did not contain extractable text")
    return ExtractedResume(
        filename=filename,
        content_type=content_type,
        text=normalized,
        character_count=len(normalized),
    )


def _extract_pdf_text(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise ResumeExtractionError("Could not parse PDF resume") from exc


def _extract_docx_text(content: bytes) -> str:
    try:
        with ZipFile(BytesIO(content)) as archive:
            _validate_docx_archive(archive)
            if "word/document.xml" not in archive.namelist():
                raise ResumeExtractionError("Unsupported resume file type")
        doc = Document(BytesIO(content))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    except ResumeExtractionError:
        raise
    except BadZipFile as exc:
        raise ResumeExtractionError("Unsupported resume file type") from exc
    except Exception as exc:
        raise ResumeExtractionError("Could not parse DOCX resume") from exc


def _validate_docx_archive(archive: ZipFile) -> None:
    members = [member for member in archive.infolist() if not member.is_dir()]
    if len(members) > MAX_DOCX_ARCHIVE_MEMBERS:
        raise ResumeExtractionError("DOCX archive contains too many files")

    total_uncompressed_bytes = 0
    for member in members:
        _validate_docx_member_name(member.filename)
        if member.flag_bits & 0x1:
            raise ResumeExtractionError("Encrypted DOCX archives are not supported")
        if member.compress_size == 0 and member.file_size > 0:
            raise ResumeExtractionError("DOCX archive has an invalid compressed member")
        if member.compress_size and member.file_size > member.compress_size * MAX_DOCX_COMPRESSION_RATIO:
            raise ResumeExtractionError("DOCX archive member exceeds the safe compression ratio")
        total_uncompressed_bytes += member.file_size
        if total_uncompressed_bytes > MAX_DOCX_UNCOMPRESSED_BYTES:
            raise ResumeExtractionError("DOCX archive is too large after decompression")


def _validate_docx_member_name(filename: str) -> None:
    path = PurePosixPath(filename.replace("\\", "/"))
    if path.is_absolute() or ".." in path.parts:
        raise ResumeExtractionError("DOCX archive contains an unsafe file path")


def _normalize_extracted_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()
